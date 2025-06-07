from docx import Document
from docx.shared import Pt
from database.db_connection import get_db_connection


def generate_word_report(request_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Get request details
    cursor.execute("""
        SELECT r.*, u.full_name as user_name 
        FROM requests r
        JOIN users u ON r.user_id = u.id
        WHERE r.id = %s
    """, (request_id,))
    request_data = cursor.fetchone()

    # Get materials for the request
    cursor.execute("""
        SELECT m.name, rm.quantity, m.unit
        FROM request_materials rm
        JOIN materials m ON rm.material_id = m.id
        WHERE rm.request_id = %s
    """, (request_id,))
    materials = cursor.fetchall()

    # Get equipment for the request
    cursor.execute("""
        SELECT e.name, re.minutes_used
        FROM request_equipment re
        JOIN equipment e ON re.equipment_id = e.id
        WHERE re.request_id = %s
    """, (request_id,))
    equipment = cursor.fetchall()

    cursor.close()
    conn.close()

    # Create Word document
    doc = Document()

    # Add title
    title = doc.add_heading(f'Заявка №{request_id}', level=1)
    title.alignment = 1  # Center alignment

    # Add request details
    doc.add_paragraph(f'Название: {request_data["title"]}')
    doc.add_paragraph(f'Описание: {request_data["description"]}')
    doc.add_paragraph(f'Площадь: {request_data["area"]} м²')
    doc.add_paragraph(f'Создал: {request_data["user_name"]}')
    doc.add_paragraph(f'Дата создания: {request_data["created_at"].strftime("%d.%m.%Y %H:%M")}')

    # Add materials table
    doc.add_heading('Материалы', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название'
    hdr_cells[1].text = 'Количество'
    hdr_cells[2].text = 'Единица'

    # Add materials rows
    for material in materials:
        row_cells = table.add_row().cells
        row_cells[0].text = material['name']
        row_cells[1].text = str(material['quantity'])
        row_cells[2].text = material['unit']

    # Add equipment table if exists
    if equipment:
        doc.add_heading('Техника', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Время использования (мин)'

        # Add equipment rows
        for item in equipment:
            row_cells = table.add_row().cells
            row_cells[0].text = item['name']
            row_cells[1].text = str(item['minutes_used'])

    # Save document
    filename = f'request_{request_id}.docx'
    doc.save(filename)

    return filename